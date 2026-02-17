import os
from typing import List, Optional
import requests
import streamlit as st
import streamlit.components.v1 as components   # needed for auto-scroll JS
import base64
import io
from docx import Document as DocxDocument

st.set_page_config(
    page_title="WS Generator",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
    <style>
    .stAppDeployButton,
    [data-testid="stStatusWidget"],
    footer { display: none; }
    .stApp header { display: none; }
    </style>
    """,
    unsafe_allow_html=True,
)

# ─────────────────────────────────────────────────────────────────────────────
# PREMIUM CSS
# ─────────────────────────────────────────────────────────────────────────────
st.markdown(
    """
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Playfair+Display:wght@400;500;600;700&display=swap" rel="stylesheet">

    <style>
      :root {
        --stone-50:     #fafaf9;
        --stone-100:    #f5f5f4;
        --stone-200:    #e7e5e4;
        --stone-300:    #d6d3d1;
        --stone-400:    #a8a29e;
        --stone-500:    #78716c;
        --stone-600:    #57534e;
        --stone-700:    #44403c;
        --stone-800:    #292524;
        --stone-900:    #1c1917;
        --espresso:     #2c2420;
        --espresso-lt:  #3d342f;
        --clay:         #9c8578;
        --cream:        #fef9f5;
        --warm-white:   #fffcf8;
        --accent:       #d97706;
        --accent-dark:  #b45309;
        --accent-light: #f59e0b;
        --accent-glow:  rgba(217,119,6,0.15);
        --accent-subtle:rgba(217,119,6,0.08);
        --success:      #059669;
        --warning:      #ea580c;
        --error:        #dc2626;
        --card:         rgba(255,252,248,0.90);
        --card-solid:   var(--warm-white);
        --glass:        rgba(255,252,248,0.70);
        --text:         var(--espresso);
        --text-muted:   var(--stone-600);
        --text-faint:   var(--clay);
        --border:       rgba(68,64,60,0.08);
        --border-med:   rgba(68,64,60,0.15);
        --shadow-sm:    0 1px 3px rgba(44,36,32,0.08), 0 1px 2px rgba(44,36,32,0.05);
        --shadow-md:    0 4px 8px rgba(44,36,32,0.10), 0 2px 4px rgba(44,36,32,0.06);
        --shadow-lg:    0 10px 20px rgba(44,36,32,0.12), 0 4px 8px rgba(44,36,32,0.08);
        --shadow-xl:    0 20px 30px rgba(44,36,32,0.14), 0 8px 12px rgba(44,36,32,0.10);
        --radius-sm:    8px;
        --radius-md:    12px;
        --radius-lg:    16px;
        --radius-xl:    24px;
      }

      *, *::before, *::after { box-sizing: border-box; }

      .stApp {
        background: linear-gradient(135deg, var(--cream) 0%, var(--stone-50) 100%);
        background-attachment: fixed;
        font-family: 'Inter', system-ui, -apple-system, sans-serif;
        color: var(--text);
        min-height: 100vh;
      }
      .block-container {
        padding-top: 0 !important;
        padding-bottom: 4rem !important;
        max-width: 1400px !important;
      }

      /* ── Typography ─── */
      h1, h2, h3 {
        font-family: 'Playfair Display', Georgia, serif;
        color: var(--espresso);
        letter-spacing: -0.025em;
        font-weight: 600;
      }
      h1 { font-size: clamp(2.8rem, 5vw, 4.2rem); line-height: 1.05; margin-bottom: 0.5rem; }
      h2 { font-size: 1.65rem; margin-bottom: 0.6rem; }
      h3 { font-size: 1.3rem; margin-bottom: 0.4rem; }

      /* ── Nav ─── */
      .ws-nav {
        position: sticky; top: 0; z-index: 999;
        backdrop-filter: blur(24px) saturate(160%);
        -webkit-backdrop-filter: blur(24px) saturate(160%);
        background: rgba(255,252,248,0.88);
        border-bottom: 1px solid var(--border);
        padding: 1rem 2.5rem;
        display: flex; align-items: center; gap: 0.8rem;
        margin: 0 -2rem 2.5rem -2rem;
        animation: fadeSlideDown 0.6s cubic-bezier(0.22,1,0.36,1) both;
        box-shadow: 0 2px 6px rgba(44,36,32,0.04);
      }
      .ws-nav-logo {
        font-family: 'Playfair Display', serif;
        font-size: 1.35rem; font-weight: 700; letter-spacing: -0.02em;
        background: linear-gradient(135deg, var(--accent) 0%, var(--accent-dark) 100%);
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        background-clip: text; margin-right: 2rem; white-space: nowrap;
      }
      .ws-nav-tab {
        font-size: 0.825rem; font-weight: 600; letter-spacing: 0.02em;
        padding: 0.5rem 1.2rem; border-radius: 999px;
        border: 1px solid transparent; color: var(--text-muted);
        cursor: pointer; transition: all 0.25s cubic-bezier(0.4,0,0.2,1);
        white-space: nowrap; text-decoration: none;
      }
      .ws-nav-tab.active {
        background: linear-gradient(135deg, var(--accent) 0%, var(--accent-dark) 100%);
        color: var(--warm-white); border-color: var(--accent);
        box-shadow: 0 3px 10px var(--accent-glow);
      }
      .ws-nav-tab:not(.active):hover {
        background: var(--accent-subtle); color: var(--accent-dark);
        border-color: var(--border-med); transform: translateY(-1px);
      }
      .ws-nav-spacer { flex: 1; }

      /* ── Hero ─── */
      .ws-hero {
        padding: 4rem 0 3rem;
        animation: fadeUp 0.7s cubic-bezier(0.22,1,0.36,1) both;
      }
      .ws-hero-tagline {
        font-family: 'Inter', sans-serif;
        font-size: 0.825rem; font-weight: 700; letter-spacing: 0.18em;
        text-transform: uppercase;
        background: linear-gradient(135deg, var(--accent) 0%, var(--accent-dark) 100%);
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        background-clip: text; margin-bottom: 1.2rem;
        display: flex; align-items: center; gap: 0.75rem;
      }
      .ws-hero-tagline::before {
        content: ''; display: inline-block; width: 40px; height: 2px;
        background: linear-gradient(90deg, var(--accent) 0%, var(--accent-dark) 100%);
        opacity: 0.9;
      }
      .ws-hero-subtitle {
        font-family: 'Inter', sans-serif; font-size: 1.08rem;
        color: var(--text-muted); font-weight: 400;
        max-width: 620px; line-height: 1.75; margin-top: 1.2rem;
      }
      .ws-badges { display: flex; flex-wrap: wrap; gap: 0.7rem; margin-top: 1.8rem; }
      .ws-badge {
        font-size: 0.75rem; font-weight: 700; letter-spacing: 0.05em;
        padding: 0.5rem 1.1rem; border-radius: 999px;
        display: inline-flex; align-items: center; gap: 0.4rem;
        box-shadow: var(--shadow-sm); transition: all 0.2s ease;
        background: linear-gradient(135deg, rgba(220,38,38,0.12) 0%, rgba(220,38,38,0.06) 100%);
        border: 1px solid rgba(220,38,38,0.25); color: #991b1b;
      }
      .ws-badge::before { content: '●'; font-size: 0.65rem; color: #dc2626; }
      .ws-badge:hover { transform: translateY(-1px); box-shadow: var(--shadow-md); }

      /* ── Steps ─── */
      .ws-steps {
        display: flex; align-items: center; justify-content: center;
        gap: 0; margin: 2rem 0 2.5rem;
        background: var(--card); border: 1px solid var(--border);
        border-radius: var(--radius-xl); padding: 1.5rem 2rem;
        box-shadow: var(--shadow-md); overflow-x: auto;
        animation: fadeUp 0.8s 0.15s cubic-bezier(0.22,1,0.36,1) both;
        backdrop-filter: blur(16px);
      }
      .ws-step {
        display: flex; align-items: center; gap: 0.6rem;
        padding: 0.6rem 1.2rem; font-size: 0.9rem; font-weight: 600;
        color: var(--text-faint); white-space: nowrap; transition: all 0.3s ease;
      }
      .ws-step-num {
        width: 30px; height: 30px; border-radius: 50%;
        display: flex; align-items: center; justify-content: center;
        font-size: 0.8rem; font-weight: 700;
        background: var(--stone-200); color: var(--stone-500);
        transition: all 0.3s ease; box-shadow: inset 0 1px 2px rgba(0,0,0,0.05);
      }
      .ws-step.active .ws-step-num {
        background: linear-gradient(135deg, var(--accent) 0%, var(--accent-dark) 100%);
        color: white; box-shadow: 0 2px 8px var(--accent-glow);
      }
      .ws-step.active { color: var(--espresso); }
      .ws-step.done .ws-step-num { background: var(--success); color: white; }
      .ws-step.done { color: var(--success); }
      .ws-step-arrow { margin: 0 0.8rem; color: var(--border-med); font-weight: 300; }

      /* ── Cards ─── */
      .ws-card {
        background: var(--card); backdrop-filter: blur(20px);
        border: 1px solid var(--border); border-radius: var(--radius-lg);
        padding: 2rem; margin-bottom: 1.5rem; box-shadow: var(--shadow-md);
        animation: fadeUp 0.6s ease both;
        transition: all 0.3s cubic-bezier(0.4,0,0.2,1);
      }
      .ws-card:hover { transform: translateY(-2px); box-shadow: var(--shadow-lg); border-color: var(--border-med); }
      .ws-card-header {
        display: flex; align-items: center; justify-content: space-between;
        margin-bottom: 1.5rem; padding-bottom: 1rem; border-bottom: 1px solid var(--border);
      }
      .ws-card-title {
        font-family: 'Playfair Display', serif; font-size: 1.4rem;
        font-weight: 600; color: var(--espresso);
        display: flex; align-items: center; gap: 0.6rem;
      }
      .ws-card-title::before {
        content: ''; width: 4px; height: 24px;
        background: linear-gradient(180deg, var(--accent) 0%, var(--accent-dark) 100%);
        border-radius: 4px;
      }
      .ws-card-badge {
        font-size: 0.7rem; font-weight: 700; letter-spacing: 0.05em;
        text-transform: uppercase; padding: 0.35rem 0.8rem; border-radius: 999px;
        background: rgba(220,38,38,0.12); color: #991b1b;
        border: 1px solid rgba(220,38,38,0.25);
      }

      /* ── FIX 2: File Uploader — Browse button must stay white ──────────────
         Rule order matters: broad color reset first, then carve-outs.
         The Browse / × button lives inside section > div > button.
         We must NOT let the wildcard `*` rule bleed into it.
      ────────────────────────────────────────────────────────────────────── */

      /* Labels above uploader: dark text ✓ */
      [data-testid="stFileUploader"] > label,
      [data-testid="stFileUploader"] > label span {
        color: var(--espresso) !important;
        font-weight: 600 !important;
        font-size: 0.95rem !important;
      }

      /* Drop-zone section: dashed warm border */
      [data-testid="stFileUploader"] section {
        background: rgba(255,252,248,0.55) !important;
        border: 2px dashed rgba(68,64,60,0.18) !important;
        border-radius: var(--radius-md) !important;
        transition: border-color 0.25s, background 0.25s, box-shadow 0.25s !important;
      }
      [data-testid="stFileUploader"] section:hover {
        border-color: rgba(217,119,6,0.50) !important;
        background: rgba(217,119,6,0.04) !important;
        box-shadow: 0 0 0 3px var(--accent-glow) !important;
      }

      /* Helper text / "Drag and drop…" span inside drop zone */
      [data-testid="stFileUploader"] section span,
      [data-testid="stFileUploader"] section p,
      [data-testid="stFileUploader"] section small {
        color: var(--stone-500) !important;
      }

      /* The Browse / Upload button — gradient amber, white text */
      [data-testid="stFileUploader"] section button,
      [data-testid="stFileUploader"] section button * {
        background: linear-gradient(135deg, var(--accent) 0%, var(--accent-dark) 100%) !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: var(--radius-sm) !important;
        font-weight: 600 !important;
        box-shadow: 0 2px 8px var(--accent-glow) !important;
      }
      [data-testid="stFileUploader"] section button:hover {
        filter: brightness(1.08) !important;
        box-shadow: 0 4px 14px var(--accent-glow) !important;
      }

      /* Uploaded file chip (the ✕ tag after upload) */
      [data-testid="stFileUploader"] [data-testid="stFileUploaderFile"],
      [data-testid="stFileUploader"] [data-testid="stFileUploaderFile"] * {
        color: var(--espresso) !important;
      }
      /* The small × delete button on the chip stays visible */
      [data-testid="stFileUploader"] [data-testid="stFileUploaderDeleteBtn"] button,
      [data-testid="stFileUploader"] [data-testid="stFileUploaderDeleteBtn"] button * {
        color: var(--stone-500) !important;
        background: transparent !important;
        box-shadow: none !important;
      }

      /* ── File Preview Strip ─── */
      .file-preview {
        margin-top: 0.75rem; background: rgba(255,252,248,0.6);
        border: 1px solid var(--border); border-radius: var(--radius-md);
        overflow: hidden; animation: fadeIn 0.3s ease;
      }
      .file-preview-row {
        display: flex; align-items: center; gap: 0.8rem;
        padding: 0.6rem 0.9rem; border-bottom: 1px solid var(--border);
        font-size: 0.85rem; transition: background 0.15s;
      }
      .file-preview-row:last-child { border-bottom: none; }
      .file-preview-row:hover { background: rgba(217,119,6,0.04); }
      .file-preview-icon {
        width: 32px; height: 32px; flex-shrink: 0;
        background: linear-gradient(135deg, var(--accent) 0%, var(--accent-dark) 100%);
        border-radius: var(--radius-sm); display: flex; align-items: center;
        justify-content: center; color: white; font-size: 0.72rem; font-weight: 700;
      }
      .file-preview-name { flex: 1; font-weight: 500; color: var(--espresso); overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
      .file-preview-size { color: var(--text-muted); font-size: 0.78rem; white-space: nowrap; }

      /* ── Input Preview Expander ─── */
      .preview-expander-wrap {
        margin-top: 0.5rem;
        border: 1px solid var(--border); border-radius: var(--radius-lg);
        overflow: hidden; background: var(--card);
        box-shadow: var(--shadow-sm);
      }
      .preview-section-label {
        font-size: 0.72rem; font-weight: 700; letter-spacing: 0.09em;
        text-transform: uppercase; color: var(--stone-400);
        padding: 0.75rem 1rem 0.3rem; margin: 0;
      }
      .pdf-viewer-frame {
        background: var(--stone-100); border-top: 1px solid var(--border);
        border-bottom: 1px solid var(--border); margin: 0.4rem 0;
      }

      /* ── Summary ─── */
      .ws-summary {
        background: linear-gradient(135deg, var(--card) 0%, rgba(255,252,248,0.95) 100%);
        backdrop-filter: blur(20px); border: 1px solid var(--border);
        border-radius: var(--radius-lg); padding: 1.8rem;
        box-shadow: var(--shadow-md); margin-bottom: 1.5rem;
        animation: fadeUp 0.7s 0.2s ease both; transition: all 0.3s ease;
      }
      .ws-summary:hover { transform: translateY(-2px); box-shadow: var(--shadow-lg); }
      .ws-summary-title {
        font-family: 'Playfair Display', serif; font-size: 1.3rem;
        font-weight: 600; color: var(--espresso);
        margin-bottom: 1.2rem; padding-bottom: 0.8rem;
        border-bottom: 1px solid var(--border);
      }
      .ws-summary-row {
        display: flex; justify-content: space-between; align-items: center;
        padding: 0.7rem 0; border-bottom: 1px solid rgba(68,64,60,0.04);
      }
      .ws-summary-row:last-child { border-bottom: none; }
      .ws-summary-key { font-size: 0.875rem; font-weight: 500; color: var(--text-muted); }
      .ws-summary-val {
        font-size: 0.875rem; font-weight: 600; padding: 0.3rem 0.7rem; border-radius: 6px;
      }
      .ws-summary-val.ok { color: var(--success); background: rgba(5,150,105,0.08); }
      .ws-summary-val.missing { color: var(--stone-400); background: rgba(168,162,158,0.08); }

      /* ── Warn ─── */
      .ws-warn {
        display: flex; align-items: flex-start; gap: 0.7rem;
        padding: 1rem 1.2rem; background: rgba(234,88,12,0.08);
        border: 1px solid rgba(234,88,12,0.25); border-radius: var(--radius-md);
        margin-bottom: 1rem; font-size: 0.9rem; color: #9a3412;
        animation: fadeIn 0.3s ease;
      }
      .ws-warn-icon { font-size: 1.1rem; flex-shrink: 0; }
      .ws-timeout-badge {
        margin-top: 1rem; text-align: center; font-size: 0.8rem;
        color: var(--text-muted); padding: 0.5rem;
        background: rgba(168,162,158,0.06); border-radius: var(--radius-sm);
      }

      /* ── Output ─── */
      .ws-output-box {
        background: linear-gradient(135deg, rgba(5,150,105,0.10) 0%, rgba(5,150,105,0.05) 100%);
        border: 1px solid rgba(5,150,105,0.30); border-radius: var(--radius-lg);
        padding: 2rem; margin: 2rem 0 1.5rem;
        animation: fadeIn 0.5s ease, pulse-green 2.5s infinite;
        box-shadow: 0 4px 12px rgba(5,150,105,0.15);
      }
      .ws-output-title {
        font-family: 'Playfair Display', serif; font-size: 1.6rem;
        font-weight: 600; color: #065f46; margin-bottom: 0.6rem;
      }
      .ws-output-meta { font-size: 0.95rem; color: #047857; line-height: 1.6; }

      /* ── Doc Preview Panel ─── */
      .doc-preview-container {
        background: var(--card); border: 1px solid var(--border);
        border-radius: var(--radius-lg); padding: 2rem; margin-top: 1.5rem;
        box-shadow: var(--shadow-md); animation: fadeUp 0.4s ease;
      }
      .doc-preview-header {
        font-family: 'Playfair Display', serif; font-size: 1.3rem;
        font-weight: 600; color: var(--espresso);
        margin-bottom: 1.5rem; padding-bottom: 1rem;
        border-bottom: 1px solid var(--border);
      }

      /* ── Buttons ─── */
      .stButton > button {
        border-radius: var(--radius-md) !important; font-weight: 600 !important;
        letter-spacing: 0.02em !important; transition: all 0.25s cubic-bezier(0.4,0,0.2,1) !important;
        border: 1px solid transparent !important; box-shadow: var(--shadow-sm) !important;
        position: relative !important; overflow: hidden !important;
      }
      .stButton > button:hover { transform: translateY(-2px) !important; box-shadow: var(--shadow-md) !important; }
      .stButton > button:active { transform: translateY(0) scale(0.99) !important; }
      .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, var(--accent) 0%, var(--accent-dark) 100%) !important;
        color: white !important;
      }
      .stButton > button[kind="primary"]:hover { box-shadow: 0 6px 16px var(--accent-glow) !important; filter: brightness(1.05) !important; }
      .stButton > button[kind="secondary"] {
        background: var(--card-solid) !important; color: var(--espresso) !important;
        border-color: var(--border-med) !important;
      }
      /* Ripple */
      .stButton > button::after {
        content: ''; position: absolute; top: 50%; left: 50%;
        width: 20px; height: 20px; background: rgba(255,255,255,0.5);
        opacity: 0; border-radius: 50%; transform: translate(-50%,-50%) scale(0);
      }
      .stButton > button:active::after { animation: ripple 0.55s ease-out; }

      /* ── Download button ─── */
      .stDownloadButton > button {
        background: linear-gradient(135deg, var(--success) 0%, #047857 100%) !important;
        color: white !important; border: none !important;
        border-radius: var(--radius-md) !important; font-weight: 600 !important;
        box-shadow: 0 4px 12px rgba(5,150,105,0.22) !important;
        transition: all 0.25s ease !important;
      }
      .stDownloadButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 22px rgba(5,150,105,0.32) !important;
        filter: brightness(1.06) !important;
      }

      /* ── Text input ─── */
      [data-testid="stTextInput"] label { font-weight: 600 !important; color: var(--espresso) !important; font-size: 0.95rem !important; }
      [data-testid="stTextInput"] input {
        background: rgba(255,252,248,0.6) !important; border: 1px solid var(--border-med) !important;
        border-radius: var(--radius-md) !important; font-size: 0.9rem !important; color: var(--espresso) !important;
        transition: border-color 0.2s, box-shadow 0.2s !important;
      }
      [data-testid="stTextInput"] input:focus {
        border-color: var(--accent) !important; box-shadow: 0 0 0 3px var(--accent-glow) !important;
      }

      /* ── Animations ─── */
      @keyframes fadeSlideDown { from { opacity:0; transform:translateY(-12px); } to { opacity:1; transform:translateY(0); } }
      @keyframes fadeUp        { from { opacity:0; transform:translateY(20px);  } to { opacity:1; transform:translateY(0); } }
      @keyframes fadeIn        { from { opacity:0; }                              to { opacity:1; } }
      @keyframes ripple        { to { transform:translate(-50%,-50%) scale(6); opacity:0; } }
      @keyframes pulse-green {
        0%,100% { box-shadow: 0 4px 12px rgba(5,150,105,0.15); }
        50%      { box-shadow: 0 4px 22px rgba(5,150,105,0.28); }
      }

      /* ── Scrollbar ─── */
      ::-webkit-scrollbar { width: 8px; height: 8px; }
      ::-webkit-scrollbar-track { background: var(--stone-100); }
      ::-webkit-scrollbar-thumb {
        background: linear-gradient(180deg, var(--accent) 0%, var(--accent-dark) 100%);
        border-radius: 4px;
      }

      /* ── Responsive ─── */
      @media (max-width: 768px) {
        .ws-nav { padding: 0.8rem 1.5rem; flex-wrap: wrap; }
        .ws-nav-logo { font-size: 1.1rem; margin-right: 1rem; }
        .ws-nav-tab { font-size: 0.75rem; padding: 0.4rem 0.9rem; }
        .ws-hero { padding: 2.5rem 0 2rem; }
        h1 { font-size: 2.2rem !important; }
        .ws-steps { padding: 1rem 1.2rem; }
        .ws-step { font-size: 0.8rem; padding: 0.5rem 0.8rem; }
        .ws-card { padding: 1.5rem; }
      }

      /* ── Focus accessibility ─── */
      button:focus-visible, input:focus-visible { outline: 2px solid var(--accent) !important; outline-offset: 2px !important; }
    </style>
    """,
    unsafe_allow_html=True,
)

# ─────────────────────────────────────────────────────────────────────────────
# Session State
# ─────────────────────────────────────────────────────────────────────────────
for k, v in {
    "clear_key":      0,
    "last_generated": None,
    "show_preview":   False,
    "scroll_to_out":  False,   # FIX 1: auto-scroll flag
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

ck = st.session_state.clear_key

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def _as_file_tuple(uf):
    return (uf.name, uf.getvalue(), uf.type or "application/octet-stream")

def _fmt_size(b: int) -> str:
    for u in ("B", "KB", "MB", "GB"):
        if b < 1024:
            return f"{b:.1f} {u}"
        b /= 1024
    return f"{b:.1f} TB"

def _ext(name: str) -> str:
    return name.rsplit(".", 1)[-1].upper() if "." in name else "FILE"

def _render_file_preview(files):
    """Compact metadata strip shown immediately after an uploader."""
    if not files:
        return
    items = files if isinstance(files, list) else [files]
    rows = "".join(
        f'<div class="file-preview-row">'
        f'  <div class="file-preview-icon">{_ext(f.name)[:3]}</div>'
        f'  <div class="file-preview-name" title="{f.name}">{f.name}</div>'
        f'  <div class="file-preview-size">{_fmt_size(len(f.getvalue()))}</div>'
        f'</div>'
        for f in items
    )
    st.markdown(f'<div class="file-preview">{rows}</div>', unsafe_allow_html=True)

# FIX 3: Full input file preview helpers
def _preview_pdf_inline(uf, height: int = 500):
    """Embed PDF as base64 iframe."""
    b64 = base64.b64encode(uf.getvalue()).decode()
    st.markdown(
        f"""
        <div class="pdf-viewer-frame">
          <iframe
            src="data:application/pdf;base64,{b64}"
            width="100%" height="{height}"
            style="border:none; display:block;"
            title="{uf.name}">
          </iframe>
        </div>
        """,
        unsafe_allow_html=True,
    )

def _preview_docx(uf):
    """Extract DOCX paragraphs and show in text area."""
    try:
        doc = DocxDocument(io.BytesIO(uf.getvalue()))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        text = f"[Could not parse DOCX: {e}]"
    st.text_area("", value=text, height=320,
                 label_visibility="collapsed", key=f"prev_docx_{uf.name}_{ck}")

def _preview_txt(uf):
    text = uf.getvalue().decode("utf-8", errors="ignore")
    st.text_area("", value=text, height=260,
                 label_visibility="collapsed", key=f"prev_txt_{uf.name}_{ck}")

def _preview_images(img_files):
    if not img_files:
        return
    n = min(len(img_files), 4)
    cols = st.columns(n)
    for i, img in enumerate(img_files):
        with cols[i % n]:
            st.image(img.getvalue(), caption=img.name, use_container_width=True)

def _render_steps(has_required: bool, generated: bool, success: bool):
    s1 = "done" if has_required else "active"
    s2 = "done" if generated else ("active" if has_required else "")
    s3 = "done" if success else ("active" if generated else "")
    s4 = "active" if success else ""
    st.markdown(
        f"""
        <div class="ws-steps">
          <div class="ws-step {s1}"><div class="ws-step-num">{"✓" if has_required else "1"}</div><span>Upload</span></div>
          <div class="ws-step-arrow">→</div>
          <div class="ws-step {s2}"><div class="ws-step-num">{"✓" if generated else "2"}</div><span>Configure</span></div>
          <div class="ws-step-arrow">→</div>
          <div class="ws-step {s3}"><div class="ws-step-num">{"✓" if success else "3"}</div><span>Generate</span></div>
          <div class="ws-step-arrow">→</div>
          <div class="ws-step {s4}"><div class="ws-step-num">{"✓" if success else "4"}</div><span>Download</span></div>
        </div>
        """,
        unsafe_allow_html=True,
    )

def extract_docx_text(content: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        return f"Error extracting text: {str(e)}"

# ─────────────────────────────────────────────────────────────────────────────
# Nav
# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
# Hero
# ─────────────────────────────────────────────────────────────────────────────
st.markdown(
    """
    <div class="ws-hero">
        <div class="ws-hero-tagline">Patent Office Document Generator</div>
        <h1>Written Submission Generator</h1>
        <div class="ws-hero-subtitle">
            Transform your patent documents into professional written submissions Upload your files and generate
            publication-ready documents in seconds.
        </div>
        <div class="ws-badges">
            <div class="ws-badge">6 Required Fields</div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

# ─────────────────────────────────────────────────────────────────────────────
# Upload Section
# ─────────────────────────────────────────────────────────────────────────────
st.markdown('<div id="upload"></div>', unsafe_allow_html=True)
colA, colB = st.columns([1.4, 1], gap="large")

with colA:
    st.markdown('<div class="ws-card">', unsafe_allow_html=True)
    st.markdown(
        """
        <div class="ws-card-header">
            <span class="ws-card-title">Required Documents</span>
            <span class="ws-card-badge">All Mandatory</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    city = st.text_input(
        "🏛️ To Office",
        value="Chennai",
        help="Enter the Patent Office city (e.g., Chennai, Mumbai, Delhi, Kolkata)",
        key=f"city_{ck}",
    )

    fer = st.file_uploader("📄 First Examination Report (FER)", type=["pdf"],
                           help="Upload the FER PDF from the patent office", key=f"fer_{ck}")
    if fer:
        _render_file_preview(fer)

    hn = st.file_uploader("📋 Hearing Notice (HN)", type=["pdf"],
                          help="Upload the Hearing Notice PDF", key=f"hn_{ck}")
    if hn:
        _render_file_preview(hn)

    spec = st.file_uploader("📑 Complete Specification", type=["pdf"],
                            help="Upload the complete specification PDF", key=f"spec_{ck}")
    if spec:
        _render_file_preview(spec)

    amended = st.file_uploader("📝 Amended Claims", type=["pdf", "docx", "txt"],
                               help="Upload amended claims (PDF, DOCX, or TXT)", key=f"amended_{ck}")
    if amended:
        _render_file_preview(amended)

    tech_imgs = st.file_uploader("🖼️ Technical Solution Diagrams",
                                 type=["png", "jpg", "jpeg"],
                                 accept_multiple_files=True,
                                 help="Upload technical diagrams (PNG/JPG, multiple allowed)",
                                 key=f"tech_imgs_{ck}")
    if tech_imgs:
        _render_file_preview(tech_imgs)

    st.markdown("</div>", unsafe_allow_html=True)

    # ── FIX 3: Input file viewer expander ────────────────────────────────────
    any_uploaded = any([fer, hn, spec, amended, bool(tech_imgs)])
    if any_uploaded:
        with st.expander("👁️  Preview input files", expanded=False):
            if city:
                st.info(f"📮 **To Office:** The Patent Office, {city}")

            if fer:
                st.markdown("##### 📄 First Examination Report")
                _preview_pdf_inline(fer, height=480)

            if hn:
                st.markdown("##### 📋 Hearing Notice")
                _preview_pdf_inline(hn, height=480)

            if spec:
                st.markdown("##### 📑 Complete Specification")
                _preview_pdf_inline(spec, height=500)

            if amended:
                ext_lower = amended.name.lower()
                st.markdown(f"##### 📝 Amended Claims — `{_ext(amended.name)}`")
                if ext_lower.endswith(".pdf"):
                    _preview_pdf_inline(amended, height=420)
                elif ext_lower.endswith(".docx"):
                    _preview_docx(amended)
                else:
                    _preview_txt(amended)

            if tech_imgs:
                st.markdown(f"##### 🖼️ Technical Diagrams ({len(tech_imgs)} file{'s' if len(tech_imgs)!=1 else ''})")
                _preview_images(tech_imgs)

    # Clear button
    if any_uploaded:
        st.markdown("<div style='margin-top:0.4rem'></div>", unsafe_allow_html=True)
        if st.button("🔄 Clear All Uploads", use_container_width=True,
                     type="secondary", key="clear_btn"):
            st.session_state.clear_key += 1
            st.session_state.show_preview = False
            st.session_state.last_generated = None
            st.rerun()

# ── Right column ─────────────────────────────────────────────────────────────
with colB:
    imgs_count = len(tech_imgs) if tech_imgs else 0
    st.markdown(
        f"""
        <div class="ws-summary">
          <div class="ws-summary-title">Summary</div>
          <div class="ws-summary-row">
            <span class="ws-summary-key">To Office</span>
            <span class="ws-summary-val {'ok' if city else 'missing'}">{'✓ ' + city if city else '✗ Missing'}</span>
          </div>
          <div class="ws-summary-row">
            <span class="ws-summary-key">FER</span>
            <span class="ws-summary-val {'ok' if fer else 'missing'}">{'✓ ' + fer.name[:20] if fer else '✗ Missing'}</span>
          </div>
          <div class="ws-summary-row">
            <span class="ws-summary-key">Hearing Notice</span>
            <span class="ws-summary-val {'ok' if hn else 'missing'}">{'✓ ' + hn.name[:20] if hn else '✗ Missing'}</span>
          </div>
          <div class="ws-summary-row">
            <span class="ws-summary-key">Specification</span>
            <span class="ws-summary-val {'ok' if spec else 'missing'}">{'✓ ' + spec.name[:20] if spec else '✗ Missing'}</span>
          </div>
          <div class="ws-summary-row">
            <span class="ws-summary-key">Amended Claims</span>
            <span class="ws-summary-val {'ok' if amended else 'missing'}">{'✓ ' + amended.name[:20] if amended else '✗ Missing'}</span>
          </div>
          <div class="ws-summary-row">
            <span class="ws-summary-key">Diagrams</span>
            <span class="ws-summary-val {'ok' if imgs_count else 'missing'}">
              {'✓ ' + str(imgs_count) + ' file' + ('s' if imgs_count != 1 else '') if imgs_count else '✗ Missing'}
            </span>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    backend_url = os.environ.get("WS_BACKEND_URL", "http://127.0.0.1:8000")
    timeout_s = 1000

    st.markdown('<div id="generate"></div>', unsafe_allow_html=True)
    st.markdown('<div class="ws-card">', unsafe_allow_html=True)
    st.markdown(
        """
        <div class="ws-card-header">
          <span class="ws-card-title">Generate</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    has_required = all([city, fer, hn, spec, amended, tech_imgs and len(tech_imgs) > 0])
    if not has_required:
        st.markdown(
    """
    <div class="ws-warn">
      <span class="ws-warn-icon">⚠</span>
      <div>
        <b>All fields are required.</b> Complete all uploads and enter the office city.
        Once you click <b>Generate Written Submission</b>, the generated document
        and download option will appear at the <b>bottom of this page</b>.
      </div>
    </div>
    """,
    unsafe_allow_html=True,
)


    st.markdown(
    """
    <div class="ws-warn" style="margin-top:0.8rem;">
      <span class="ws-warn-icon">⚠</span>
      <div>
        <b>Desclaimer:</b><br/>
        Prior art details (D1, D2, etc.), disclosures, and related mappings are automatically extracted from the FER.
        Due to variations in IPO formatting (e.g., combined references on a single line or unconventional phrasing),
        the extraction may not always be fully accurate.<br/>
        Please carefully verify all prior art references and related content before finalizing the Written Submission.
      </div>
    </div>
    """,
    unsafe_allow_html=True,
)


    go = st.button(
        "⚡  Generate Written Submission",
        type="primary",
        use_container_width=True,
        disabled=not has_required,
    )

    st.markdown(
        f'<div class="ws-timeout-badge">⏱ Request Timeout: {timeout_s}s</div>',
        unsafe_allow_html=True,
    )
    st.markdown("</div>", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# Step Indicator
# ─────────────────────────────────────────────────────────────────────────────
_render_steps(
    has_required=has_required,
    generated=go,
    success=st.session_state.last_generated is not None,
)

# ─────────────────────────────────────────────────────────────────────────────
# GENERATION LOGIC — pipeline unchanged
# ─────────────────────────────────────────────────────────────────────────────
if go:
    missing = []
    if not city:        missing.append("To Office (City)")
    if not fer:         missing.append("FER")
    if not hn:          missing.append("Hearing Notice")
    if not spec:        missing.append("Specification")
    if not amended:     missing.append("Amended Claims")
    if not (tech_imgs and len(tech_imgs) > 0):
        missing.append("Technical Diagrams")

    if missing:
        st.error("Please upload all required files: " + ", ".join(missing))
        st.stop()

    endpoint = backend_url.rstrip("/") + "/api/generate"

    files_list = [
        ("fer",            _as_file_tuple(fer)),
        ("hn",             _as_file_tuple(hn)),
        ("specification",  _as_file_tuple(spec)),
        ("amended_claims", _as_file_tuple(amended)),
    ]
    for img in tech_imgs:
        files_list.append(("tech_solution_images", _as_file_tuple(img)))

    data_dict = {"city": city}

    with st.status("Generating…", expanded=False) as status:
        try:
            st.write("📡 Calling:", endpoint)
            r = requests.post(endpoint, files=files_list, data=data_dict, timeout=timeout_s)
            st.write("📊 Status:", r.status_code)
            if r.status_code != 200:
                st.code(r.text[:4000])
                status.update(label="❌ Failed", state="error")
                st.stop()

            filename = "written_submission.docx"
            cd = r.headers.get("content-disposition", "")
            if "filename=" in cd:
                fn = cd.split("filename=")[-1].strip().strip('"')
                if fn:
                    filename = fn

            status.update(label="✅ Done", state="complete")
            st.session_state.last_generated = (filename, r.content)
            st.session_state.scroll_to_out = True   # FIX 1: arm the scroll flag

        except requests.exceptions.RequestException as e:
            status.update(label="❌ Failed", state="error")
            st.error(f"Request error: {e}")

    # FIX 1: Rerun immediately so output section renders and scroll JS fires
    if st.session_state.last_generated is not None:
        st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
# Output anchor + FIX 1: Auto-scroll JS
# ─────────────────────────────────────────────────────────────────────────────
st.markdown('<div id="output"></div>', unsafe_allow_html=True)

# Fire scroll once, then reset flag so it doesn't re-trigger on every rerun
if st.session_state.scroll_to_out:
    components.html(
        """
        <script>
          (function() {
            try {
              var doc = window.parent.document;
              var target = doc.getElementById("output");
              if (target) {
                target.scrollIntoView({ behavior: "smooth", block: "start" });
              } else {
                window.parent.scrollTo({
                  top: window.parent.document.body.scrollHeight,
                  behavior: "smooth"
                });
              }
            } catch(e) {}
          })();
        </script>
        """,
        height=0,
        scrolling=False,
    )
    st.session_state.scroll_to_out = False  # reset — won't re-fire on next interaction

# ─────────────────────────────────────────────────────────────────────────────
# Output Download & Preview Section
# ─────────────────────────────────────────────────────────────────────────────
if st.session_state.last_generated is not None:
    filename, content = st.session_state.last_generated

    st.markdown(
        """
        <div class="ws-output-box">
          <div class="ws-output-title">✓  Written Submission Ready</div>
          <div class="ws-output-meta">Your document has been generated and is ready for download or preview below.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    col_dl, col_preview, col_close = st.columns([1, 1, 1])

    with col_dl:
        st.download_button(
            "⬇  Download DOCX",
            data=content,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with col_preview:
        if st.button("👁  Preview Document", use_container_width=True,
                     type="secondary", key="btn_prev"):
            st.session_state.show_preview = not st.session_state.show_preview
            st.rerun()

    with col_close:
        if st.session_state.show_preview:
            if st.button("✖  Close Preview", use_container_width=True,
                         type="secondary", key="btn_close"):
                st.session_state.show_preview = False
                st.rerun()

    if st.session_state.show_preview:
        st.markdown(
            """
            <div class="doc-preview-container">
              <div class="doc-preview-header">📄 Document Preview</div>
            """,
            unsafe_allow_html=True,
        )

        st.info(
            f"**File:** {filename}  \n"
            f"**Size:** {_fmt_size(len(content))}  \n"
            f"**Office:** The Patent Office, {city}"
        )

        with st.spinner("Extracting document content…"):
            doc_text = extract_docx_text(content)

        st.text_area("", value=doc_text, height=420,
                     label_visibility="collapsed", key="out_preview_text")

        st.markdown(
            f"""
            <div style="background:rgba(217,119,6,0.06); padding:1.2rem; border-radius:12px;
                        margin-top:1.2rem; border:1px solid rgba(217,119,6,0.15);">
              <strong style="color:var(--accent-dark); font-size:0.95rem;">Document Details</strong>
              <div style="color:var(--text-muted); font-size:0.875rem; margin-top:0.6rem; line-height:1.8;">
                • Format: Microsoft Word Document (.docx)<br>
                • Status: Ready for download<br>
                • Office: The Patent Office, {city}<br>
                • Characters: {len(doc_text):,}
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown("</div>", unsafe_allow_html=True)

    st.success(f"✅ Written Submission generated → `{filename}`")
