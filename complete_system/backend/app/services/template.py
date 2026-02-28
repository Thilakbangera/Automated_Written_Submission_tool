from __future__ import annotations

import re
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.table import Table
from docx.text.paragraph import Paragraph


def _replace_runs_in_paragraph(p: Paragraph, mapping: Dict[str, str]) -> None:
    for run in p.runs:
        for k, v in mapping.items():
            if k in run.text:
                run.text = run.text.replace(k, v)


def _iter_paragraphs_in_doc(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _replace_block_placeholder_with_paragraphs(doc: Document, placeholder: str, value: str) -> None:
    """Replace a placeholder paragraph with multiple paragraphs split by blank lines."""
    blocks: List[str] = [b.strip() for b in re.split(r"\n{2,}", value or "") if b and b.strip()]
    for p in list(_iter_paragraphs_in_doc(doc)):
        full = p.text or ""
        if placeholder not in full:
            continue

        if not blocks:
            p.text = full.replace(placeholder, "").strip()
            return

        from docx.oxml import OxmlElement

        style = p.style
        align = p.alignment
        before, after = full.split(placeholder, 1)

        first_text = (before or "").strip()
        if first_text:
            first_text = f"{first_text}\n{blocks[0]}"
        else:
            first_text = blocks[0]

        p.text = first_text
        p.style = style
        p.alignment = align

        anchor = p
        for block in blocks[1:]:
            el = OxmlElement("w:p")
            anchor._p.addnext(el)
            np = Paragraph(el, anchor._parent)
            np.style = style
            np.alignment = align
            np.text = block
            anchor = np

        tail = (after or "").strip()
        if tail:
            el = OxmlElement("w:p")
            anchor._p.addnext(el)
            np = Paragraph(el, anchor._parent)
            np.style = style
            np.alignment = align
            np.text = tail
        return


def _insert_table_after(paragraph: Paragraph, rows: int, cols: int) -> Table:
    """Insert a table *after* the given paragraph."""
    tbl = paragraph._parent.add_table(rows=rows, cols=cols)  # type: ignore
    paragraph._p.addnext(tbl._tbl)  # type: ignore
    return tbl


def ensure_feature_table(doc: Document, claim1_text: str, d1d2_text: str) -> None:
    """Ensure the 'Applicant claimed feature vs D1-D2 disclosed features' table exists.

    V3 requirement: table must never be missing.
    Strategy:
      1) If template includes {{FEATURE_TABLE}} marker -> replace it with a table.
      2) Else insert the table right after the paragraph containing {{AMENDED_CLAIM_1}}.
    """
    marker_p: Optional[Paragraph] = None
    claim1_p: Optional[Paragraph] = None

    for p in doc.paragraphs:
        if "{{FEATURE_TABLE}}" in p.text:
            marker_p = p
            break
        if "{{AMENDED_CLAIM_1}}" in p.text:
            claim1_p = p

    anchor = marker_p or claim1_p
    if anchor is None:
        return

    # If the table already exists, do nothing.
    for t in doc.tables:
        try:
            if t.cell(0, 0).text.strip().lower() == "applicant claimed feature":
                return
        except Exception:
            continue

    tbl = _insert_table_after(anchor, rows=2, cols=2)
    tbl.cell(0, 0).text = "Applicant claimed feature"
    tbl.cell(0, 1).text = "D1-D2 disclosed features"
    tbl.cell(1, 0).text = claim1_text or ""
    tbl.cell(1, 1).text = d1d2_text or ""

    # Remove marker paragraph if present
    if marker_p is not None:
        marker_p.text = ""


def insert_prior_art_analysis_before_feature_table(
    doc: Document,
    analysis_text: str,
    diagram_images: Optional[list] = None,
    analysis_sequence: Optional[list] = None,
) -> None:
    """Insert prior-art analysis text/images immediately before feature table."""
    txt = (analysis_text or "").strip()
    images = diagram_images or []
    sequence = analysis_sequence or []
    if not txt and not images and not sequence:
        return

    feature_table: Optional[Table] = None
    for t in doc.tables:
        try:
            if t.cell(0, 0).text.strip().lower() == "applicant claimed feature":
                feature_table = t
                break
        except Exception:
            continue

    if feature_table is None:
        return

    from docx.oxml import OxmlElement

    def _insert_paragraph_before_table(text: str = "") -> Paragraph:
        p_el = OxmlElement("w:p")
        feature_table._tbl.addprevious(p_el)  # type: ignore[attr-defined]
        p = Paragraph(p_el, feature_table._parent)
        if text:
            p.text = text
        return p

    def _insert_text_before_table(text: str, make_red: bool = False) -> None:
        p = _insert_paragraph_before_table()
        run = p.add_run(text)
        if make_red:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    if sequence:
        for item in sequence:
            if not isinstance(item, dict):
                continue
            kind = str(item.get("kind", "")).strip().lower()
            if kind == "image":
                img_path = str(item.get("path", "")).strip()
                if not img_path:
                    continue
                img_p = _insert_paragraph_before_table()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    img_p.add_run().add_picture(img_path, width=Inches(5.8))
                except Exception:
                    continue
            else:
                text = str(item.get("text", "")).strip()
                if text:
                    is_combined_diff = bool(re.match(r"^\s*Combined\s+difference\s+over\b", text, re.I))
                    _insert_text_before_table(text=text, make_red=is_combined_diff)
        return

    if txt:
        for block in [b.strip() for b in re.split(r"\n{2,}", txt) if b and b.strip()]:
            is_combined_diff = bool(re.match(r"^\s*Combined\s+difference\s+over\b", block, re.I))
            _insert_text_before_table(text=block, make_red=is_combined_diff)

    for item in images:
        if isinstance(item, dict):
            img_path = str(item.get("path", "")).strip()
        else:
            img_path = str(item).strip()
        if not img_path:
            continue

        img_p = _insert_paragraph_before_table()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            img_p.add_run().add_picture(img_path, width=Inches(5.8))
        except Exception:
            # If image load fails, skip image and continue WS generation.
            continue


def remove_empty_claim_sections(doc: Document, max_claim_number: int) -> None:
    """Remove 'Regarding Claim N' sections for claims that don't exist.
    
    Template has static sections for claims 2-10. If actual claims only go to 7,
    we need to delete the sections for claims 8, 9, 10.
    """
    if max_claim_number >= 10:
        return  # All template sections are valid
    
    paragraphs_to_remove = []
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        # Check if this is a "Regarding Claim N:" heading
        match = re.match(r'^\s*Regarding\s+Claim\s+(\d+)\s*:\s*$', p.text, re.I)
        if match:
            claim_num = int(match.group(1))
            # If this claim number doesn't exist, mark this paragraph and the next one for removal
            if claim_num > max_claim_number:
                paragraphs_to_remove.append(i)
                # The next paragraph contains the claim text, remove it too
                if i + 1 < len(doc.paragraphs):
                    paragraphs_to_remove.append(i + 1)
                    i += 1  # Skip the next paragraph since we're removing it
        i += 1
    
    # Remove paragraphs in reverse order to maintain indices
    for idx in reversed(paragraphs_to_remove):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)


def _style_reply_by_drafter_marker(doc: Document) -> None:
    marker = "[REPLY BY DRAFTER]"
    for p in _iter_paragraphs_in_doc(doc):
        full = p.text or ""
        if marker not in full:
            continue
        parts = full.split(marker)
        p.text = ""
        for i, seg in enumerate(parts):
            if seg:
                p.add_run(seg)
            if i < len(parts) - 1:
                r = p.add_run(marker)
                r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def _is_heading_or_side_heading_line(line: str) -> bool:
    t = re.sub(r"\s+", " ", (line or "")).strip()
    if not t:
        return False
    if len(t) > 180:
        return False

    # Standalone headings and sub-headings.
    if re.match(
        r"^(?:"
        r"Applicant Submission|REPLY TO OBJECTION|STATEMENT REGARDING SUBSTANCE OF HEARING|"
        r"Formal Requirement(?:\(s\)|s)?|Clarity and Conciseness|Definitiveness|Definiteness|"
        r"Invention\s+u/s\b.*|Other Requirement(?:\(s\)|s)?|Prior Art|Novelty|Inventive Step|"
        r"NON-PATENTABILITY U/S 3|TECHNICAL ADVANCEMENT|TECHNICAL PROBLEM SOLVED BY THE INVENITON|"
        r"TECHNICAL SOLUTION SOLVED BY THE INVENITON|Technical Effect|Regarding Claim \d+|"
        r"Yours faithfully|Enclosure"
        r")\s*:?\s*$",
        t,
        re.I,
    ):
        return True

    # Pure uppercase heading lines.
    if t == t.upper() and re.search(r"[A-Z]", t):
        return True

    return False


def _style_headings_and_side_headings(doc: Document) -> None:
    marker = "[REPLY BY DRAFTER]"

    def _append_line(paragraph: Paragraph, line_text: str, is_heading: bool):
        last_run = None
        if marker not in line_text:
            run = paragraph.add_run(line_text)
            if is_heading:
                run.font.bold = True
                run.font.underline = True
            return run

        parts = line_text.split(marker)
        for i, seg in enumerate(parts):
            if seg:
                run = paragraph.add_run(seg)
                if is_heading:
                    run.font.bold = True
                    run.font.underline = True
                last_run = run
            if i < len(parts) - 1:
                mrun = paragraph.add_run(marker)
                mrun.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                last_run = mrun
        return last_run

    for p in _iter_paragraphs_in_doc(doc):
        full = p.text or ""
        if not full.strip():
            continue

        lines = full.split("\n")
        has_multiline = len(lines) > 1
        if has_multiline:
            if not any(_is_heading_or_side_heading_line(ln) for ln in lines) and marker not in full:
                continue
            p.text = ""
            for i, ln in enumerate(lines):
                is_heading = _is_heading_or_side_heading_line(ln)
                last = _append_line(p, ln, is_heading)
                if i < len(lines) - 1:
                    if last is None:
                        last = p.add_run("")
                    last.add_break()
            continue

        if _is_heading_or_side_heading_line(full):
            for run in p.runs:
                if run.text and run.text.strip():
                    run.font.bold = True
                    run.font.underline = True


def replace_placeholders(doc_path: str, out_path: str, mapping: Dict[str, str]) -> None:
    doc = Document(doc_path)

    # Ensure feature table BEFORE replacement, so anchor placeholders exist.
    # Use compact 'preamble + features' for the table (preferred), falling back to full Claim 1.
    ensure_feature_table(
        doc,
        claim1_text=mapping.get("{{AMENDED_CLAIM_1}}", "") or mapping.get("{{CLAIM1_FEATURES}}", ""),
        d1d2_text=mapping.get("{{D1D2_DISCLOSURE}}", ""),
    )
    insert_prior_art_analysis_before_feature_table(
        doc,
        analysis_text=mapping.get("__PRIOR_ARTS_ABSTRACTS_AND_DIFF__", ""),
        diagram_images=mapping.get("__PRIOR_ART_DIAGRAM_IMAGES__", []),
        analysis_sequence=mapping.get("__PRIOR_ART_SEQUENCE__", []),
    )
    _replace_block_placeholder_with_paragraphs(
        doc,
        placeholder="{{FORMAL_OBJECTIONS_REPLY}}",
        value=mapping.get("{{FORMAL_OBJECTIONS_REPLY}}", ""),
    )
    _replace_block_placeholder_with_paragraphs(
        doc,
        placeholder="{{TECH_EFFECT}}",
        value=mapping.get("{{TECH_EFFECT}}", ""),
    )
    _replace_block_placeholder_with_paragraphs(
        doc,
        placeholder="{{AMENDED_CLAIM_n}}",
        value=mapping.get("{{AMENDED_CLAIM_n}}", ""),
    )

    for p in _iter_paragraphs_in_doc(doc):
        _replace_runs_in_paragraph(p, mapping)

    _style_reply_by_drafter_marker(doc)
    _style_headings_and_side_headings(doc)

    # Remove empty claim sections for claims that don't exist
    max_claim = mapping.get("__CLAIM_MAX__", 10)
    if isinstance(max_claim, int):
        remove_empty_claim_sections(doc, max_claim)

    # Insert Technical Solution diagrams (optional)
    images = mapping.get("__TECH_SOLUTION_IMAGES__") if isinstance(mapping, dict) else None
    if images:
        _insert_images_with_captions(doc, "{{TECH_SOLUTION_IMAGES}}", images, start_fig_no=1)

    doc.save(out_path)


def insert_tech_solution_images(doc: Document, image_paths: list, heading_text: str = "TECHNICAL SOLUTION", start_fig_no: int = 1):
    """Insert centered diagram screenshots under the TECHNICAL SOLUTION heading with a page break and FIG captions."""
    if not image_paths:
        return

    heading_p = None
    for p in doc.paragraphs:
        if heading_text.upper() in (p.text or "").upper():
            heading_p = p
            break
    if heading_p is None:
        return

    new_paras = []

    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)
    new_paras.append(pb)

    fig_no = start_fig_no
    for img in image_paths:
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture(str(img), width=Inches(5.8))
        new_paras.append(img_p)

        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.add_run(f"FIG. {fig_no}")
        new_paras.append(cap_p)

        reply_p = doc.add_paragraph()
        reply_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        reply_run = reply_p.add_run("[Enter Description of the diagram]")
        reply_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        new_paras.append(reply_p)
        fig_no += 1

    anchor = heading_p._p
    for np in new_paras:
        anchor.addnext(np._p)
        anchor = np._p


def _insert_images_with_captions(doc: Document, placeholder: str, image_paths: list, start_fig_no: int = 1):
    """Replace placeholder paragraph with:
    - page break
    - centered images
    - centered captions (FIG. n)
    Inserted exactly at the placeholder location (not at document end).
    """
    if not image_paths:
        return

    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    for p in doc.paragraphs:
        if placeholder in p.text:
            # clear placeholder text
            for r in p.runs:
                r.text = ""

            # insert a page break in the placeholder paragraph itself
            p.add_run().add_break(WD_BREAK.PAGE)

            # we will insert new paragraphs immediately after placeholder paragraph
            anchor = p
            fig_no = start_fig_no

            def insert_paragraph_after(anchor_paragraph):
                new_p = OxmlElement("w:p")
                anchor_paragraph._p.addnext(new_p)
                return Paragraph(new_p, anchor_paragraph._parent)

            for img_path in image_paths:
                # image paragraph
                img_p = insert_paragraph_after(anchor)
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(str(img_path), width=Inches(5.8))

                # caption paragraph
                cap_p = insert_paragraph_after(img_p)
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.add_run(f"FIG. {fig_no}")
                fig_no += 1
                
                # drafter reply marker just after each diagram block
                reply_p = insert_paragraph_after(cap_p)
                reply_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                reply_run = reply_p.add_run("[Enter Description of the diagram]")
                reply_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

                anchor = reply_p  # continue inserting after marker

            return

    for p in doc.paragraphs:
        if placeholder in p.text:
            # Clear placeholder text
            for r in p.runs:
                r.text = ""
            # Page break before first diagram
            p.add_run().add_break(WD_BREAK.PAGE)

            fig_no = start_fig_no
            for img_path in image_paths:
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(str(img_path), width=Inches(5.8))

                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.add_run(f"FIG. {fig_no}")
                fig_no += 1

                reply_p = doc.add_paragraph()
                reply_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                reply_run = reply_p.add_run("[Enter Description of the diagram]")
                reply_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            return
