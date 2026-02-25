from __future__ import annotations

from collections import Counter
import json
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from urllib.error import HTTPError, URLError
from urllib.parse import quote
from urllib.request import urlopen

import pdfplumber

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None


@dataclass
class PriorArt:
    label: str
    docno: str
    date: str = ""


@dataclass
class CaseMeta:
    app_no: str = ""
    filed_on: str = ""
    applicant: str = ""
    controller: str = ""
    agents: str = ""
    fer_dispatch_date: str = ""   # FER email/dispatch date (top page)
    hn_dispatch_date: str = ""
    hearing_date: str = ""
    hearing_time: str = ""
    hearing_duration: str = ""
    hearing_mode: str = "Video Conferencing"
    fer_date: str = ""            # Sometimes same as fer_dispatch_date
    fer_reply_date: str = ""
    prior_arts: List[PriorArt] = None
    d1_disclosure: str = ""
    d2_disclosure: str = ""


def _normalize_pdf_line(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip().lower()


def _line_is_page_marker(s: str) -> bool:
    x = _normalize_pdf_line(s)
    if not x:
        return True
    if re.fullmatch(r"[-\u2013\u2014]?\s*\d{1,4}\s*[-\u2013\u2014]?", x):
        return True
    if re.fullmatch(r"\d+\s*/\s*\d+", x):
        return True
    if re.fullmatch(r"(?:p|pg|page)\.?\s*\d+\s*/\s*\d+", x):
        return True
    if re.fullmatch(r"page\s*\d+(\s*of\s*\d+)?", x):
        return True
    return False


def _non_ascii_ratio(s: str) -> float:
    if not s:
        return 0.0
    non_ascii = sum(1 for ch in s if ord(ch) > 127)
    return non_ascii / max(1, len(s))


def _duration_from_time_range(text: str) -> str:
    m = re.search(
        r"(\d{1,2}:\d{2})\s*(?:HRS|IST|AM|PM)?\s*(?:to|\-|–|—)\s*(\d{1,2}:\d{2})",
        text or "",
        re.I,
    )
    if not m:
        return ""
    try:
        sh, sm = [int(p) for p in m.group(1).split(":")]
        eh, em = [int(p) for p in m.group(2).split(":")]
        start = sh * 60 + sm
        end = eh * 60 + em
        if end < start:
            end += 24 * 60
        mins = end - start
        if mins <= 0:
            return ""
        if mins % 60 == 0:
            hrs = mins // 60
            return f"{hrs} hour" + ("s" if hrs != 1 else "")
        return f"{mins} minutes"
    except Exception:
        return ""


def _duration_from_phrase(text: str) -> str:
    txt = text or ""
    m = re.search(r"\bfor\s*\(?\s*([0-9]{1,3}\s*(?:minutes?|mins?|hours?|hrs?))\s*\)?", txt, re.I)
    if m:
        return _clean(m.group(1))
    m = re.search(r"\bduration\s*(?:of)?\s*[:\-]?\s*\(?\s*([0-9]{1,3}\s*(?:minutes?|mins?|hours?|hrs?))\s*\)?", txt, re.I)
    if m:
        return _clean(m.group(1))
    return ""


def _keep_even_if_repeated(norm_line: str) -> bool:
    if not norm_line:
        return False
    has_date = bool(re.search(r"\d{1,2}[./-]\d{1,2}[./-]\d{2,4}", norm_line))
    if not has_date:
        return False
    if re.fullmatch(r"(?:date\s*[:\-]?\s*)?\d{1,2}[./-]\d{1,2}[./-]\d{2,4}", norm_line, re.I):
        return True
    return bool(re.search(r"\b(date|dated|dispatch|hearing|email)\b", norm_line, re.I))


def _line_is_edge_header_footer_noise(s: str) -> bool:
    x = _normalize_pdf_line(s)
    if not x or len(x) > 140:
        return False
    return bool(
        re.search(r"\bpatent\s+agent\b", x, re.I)
        or re.search(r"\boffice\s+of\s+the\s+controller\s+general\b", x, re.I)
        or re.search(r"\bintellectual\s+property\s+india\b", x, re.I)
    )


def read_pdf_text(path: str) -> str:
    pages_lines: List[List[str]] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            lines = [ln.strip() for ln in t.splitlines() if ln and ln.strip()]
            pages_lines.append(lines)

    if not pages_lines:
        return ""

    # Remove repeated header/footer lines appearing on most pages.
    norm_page_counts: Counter[str] = Counter()
    for lines in pages_lines:
        seen = { _normalize_pdf_line(ln) for ln in lines if _normalize_pdf_line(ln) }
        for n in seen:
            norm_page_counts[n] += 1

    repeated_threshold = max(2, int(len(pages_lines) * 0.6))
    repeated_lines = {
        n
        for n, cnt in norm_page_counts.items()
        if cnt >= repeated_threshold and len(n) <= 180
    }

    cleaned_pages: List[str] = []
    for lines in pages_lines:
        cleaned: List[str] = []
        total = len(lines)
        for idx, ln in enumerate(lines):
            n = _normalize_pdf_line(ln)
            if not n:
                continue
            if n in repeated_lines and not _keep_even_if_repeated(n):
                continue
            if (idx <= 2 or idx >= total - 3) and _line_is_edge_header_footer_noise(ln):
                continue
            if _line_is_page_marker(ln):
                continue
            if "(cid:" in ln.lower():
                continue
            if _non_ascii_ratio(ln) > 0.45 and len(ln) > 8:
                continue
            cleaned.append(ln)
        cleaned_pages.append("\n".join(cleaned))

    return "\n".join([p for p in cleaned_pages if p.strip()])


def read_pdf_text_preserve_layout(path: str) -> str:
    """Read PDF text while preserving in-line spacing/indentation for claim blocks."""
    pages_lines: List[List[str]] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text(layout=True) or page.extract_text() or ""
            lines = [ln.rstrip("\r") for ln in t.splitlines() if ln and ln.strip()]
            pages_lines.append(lines)

    if not pages_lines:
        return ""

    norm_page_counts: Counter[str] = Counter()
    for lines in pages_lines:
        seen = {_normalize_pdf_line(ln) for ln in lines if _normalize_pdf_line(ln)}
        for n in seen:
            norm_page_counts[n] += 1

    repeated_threshold = max(2, int(len(pages_lines) * 0.6))
    repeated_lines = {
        n
        for n, cnt in norm_page_counts.items()
        if cnt >= repeated_threshold and len(n) <= 180
    }

    cleaned_pages: List[str] = []
    for lines in pages_lines:
        cleaned: List[str] = []
        total = len(lines)
        for idx, ln in enumerate(lines):
            n = _normalize_pdf_line(ln)
            if not n:
                continue
            if n in repeated_lines and not _keep_even_if_repeated(n):
                continue
            if (idx <= 2 or idx >= total - 3) and _line_is_edge_header_footer_noise(ln):
                continue
            if _line_is_page_marker(ln):
                continue
            if "(cid:" in ln.lower():
                continue
            if _non_ascii_ratio(ln) > 0.45 and len(ln) > 8:
                continue
            cleaned.append(ln.rstrip())
        cleaned_pages.append("\n".join(cleaned))

    return "\n".join([p for p in cleaned_pages if p.strip()])


def read_docx_text(path: str) -> str:
    if Document is None:
        return ""
    doc = Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        parts.append(p.text)
    return "\n".join(parts)


def read_text_any(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return read_pdf_text(path)
    if ext == ".docx":
        return read_docx_text(path)
    try:
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def _clean(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()


_PRIOR_ART_ABSTRACT_HEADINGS = [
    "abstract",
    "abstrait",
    "abrege",
    "abrégé",
    "resumen",
    "resumo",
    "riassunto",
    "zusammenfassung",
    "samenvatting",
    "sammanfattning",
    "özet",
    "摘要",
    "要約",
    "概要",
]

_PRIOR_ART_STOP_HEADINGS = [
    "claim",
    "claims",
    "what is claimed",
    "field",
    "technical field",
    "background",
    "detailed description",
    "description",
    "brief description",
    "embodiment",
    "drawings",
    "prior art",
    "revendications",
    "reivindicaciones",
    "ansprüche",
    "权利要求",
]

_TRANSLATE_TIMEOUT_S = 8
_TRANSLATE_MAX_CHARS = 2800
_PRIOR_ART_MAX_ABSTRACT_WORDS = 900


def _read_prior_art_pdf_lines(path: str, max_pages: int = 5) -> List[str]:
    pages_lines: List[List[str]] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[: max(1, max_pages)]:
                t = page.extract_text(layout=True) or page.extract_text() or ""
                page_lines: List[str] = []
                for raw in t.splitlines():
                    line = re.sub(r"\s+", " ", (raw or "").strip())
                    if not line:
                        page_lines.append("")
                        continue
                    if _is_prior_art_header_footer_noise(line):
                        continue
                    if "(cid:" in line.lower():
                        continue
                    page_lines.append(line)
                pages_lines.append(page_lines)
    except Exception:
        return []

    if not pages_lines:
        return []

    # Remove repeated page-edge noise (headers/footers) seen across most pages.
    norm_page_counts: Counter[str] = Counter()
    for page_lines in pages_lines:
        seen = {_normalize_pdf_line(ln) for ln in page_lines if ln and ln.strip()}
        for n in seen:
            norm_page_counts[n] += 1

    repeated_threshold = max(2, int(len(pages_lines) * 0.65))
    repeated_lines = {
        n
        for n, cnt in norm_page_counts.items()
        if cnt >= repeated_threshold and len(n) <= 180
    }

    out: List[str] = []
    prev_blank = False
    for page_lines in pages_lines:
        for ln in page_lines:
            if not ln:
                if out and not prev_blank:
                    out.append("")
                prev_blank = True
                continue
            n = _normalize_pdf_line(ln)
            if n in repeated_lines and not re.search(r"\babstract\b|^\s*\[?\s*57\s*\]?\s*abstract\b", n, re.I):
                continue
            if _is_prior_art_header_footer_noise(ln):
                continue
            out.append(ln)
            prev_blank = False
    return out


def _is_prior_art_metadata_line(line: str) -> bool:
    x = _clean(line)
    if not x:
        return True
    low = x.lower()
    if re.fullmatch(r"[0-9\W_]+", x):
        return True
    if re.fullmatch(r"[a-z]?\d{2,}[a-z0-9/\-]*", low):
        return True
    if re.search(
        r"\b(application|publication|applicant|inventor|priority|filing|date|int\.?cl|ipc|cpc|attorney|agent)\b",
        low,
    ):
        return True
    return False


def _looks_like_prior_art_heading(line: str) -> bool:
    x = _clean(line)
    if not x:
        return False
    if len(x) > 140:
        return False
    low = x.lower().rstrip(":")
    if any(h in low for h in _PRIOR_ART_STOP_HEADINGS):
        return True
    if x.endswith(":"):
        return True
    words = re.findall(r"[A-Za-z0-9\u00C0-\u024F]+", x)
    if not words:
        return False
    upper_words = sum(1 for w in words if w.upper() == w and len(w) > 1)
    return upper_words >= max(2, int(len(words) * 0.65))


def _clean_prior_art_abstract_text(s: str) -> str:
    txt = (s or "").replace("\r", "\n")
    lines = [ln.strip() for ln in txt.splitlines()]

    out: List[str] = []
    prev_blank = False
    for ln in lines:
        if not ln:
            if out and not prev_blank:
                out.append("")
            prev_blank = True
            continue
        out.append(ln)
        prev_blank = False

    txt = "\n".join(out).strip()
    paras = [p for p in re.split(r"\n{2,}", txt) if p and p.strip()]
    cleaned_paras: List[str] = []
    for p in paras:
        q = re.sub(r"-\s*\n\s*", "", p)
        q = re.sub(r"\s*\n\s*", " ", q)
        q = re.sub(r"\s+([,.;:])", r"\1", q)
        q = re.sub(r"\(\s+", "(", q)
        q = re.sub(r"\s+\)", ")", q)
        q = _clean(q)
        if q:
            cleaned_paras.append(q)
    return "\n\n".join(cleaned_paras).strip()


def _is_prior_art_header_footer_noise(line: str) -> bool:
    x = _clean(line)
    if not x:
        return True
    low = _normalize_pdf_line(x)
    if not low:
        return True
    if _line_is_page_marker(x):
        return True
    if _line_is_edge_header_footer_noise(x):
        return True
    if re.search(r"\b(?:https?://|www\.)\S+", low):
        return True
    if re.search(r"\b(?:google\s+patents|patentscope|espacenet|patent\s+images|lens\.org|wipo)\b", low):
        return True
    if re.search(r"\b(?:copyright|all rights reserved)\b", low):
        return True
    if re.fullmatch(r"(?:date\s*[:\-]\s*)?\d{1,2}[./-]\d{1,2}[./-]\d{2,4}", low):
        return True
    if re.fullmatch(r"(?:page|pg)\s*\d+\s*(?:of|/)\s*\d+", low):
        return True
    return False


def _trim_abstract_without_mid_sentence_cut(text: str, max_words: int = _PRIOR_ART_MAX_ABSTRACT_WORDS) -> str:
    txt = _clean_prior_art_abstract_text(text)
    if not txt:
        return ""
    words = txt.split()
    if len(words) <= max_words:
        return txt
    cropped = " ".join(words[:max_words]).strip()
    last_punct = max(cropped.rfind(". "), cropped.rfind("; "), cropped.rfind("? "), cropped.rfind("! "))
    if last_punct >= int(len(cropped) * 0.6):
        cropped = cropped[: last_punct + 1]
    return cropped.strip()


def _looks_non_english(text: str) -> bool:
    txt = _clean(text)
    if not txt:
        return False

    cjk_chars = sum(
        1
        for ch in txt
        if ("\u4e00" <= ch <= "\u9fff") or ("\u3040" <= ch <= "\u30ff") or ("\uac00" <= ch <= "\ud7af")
    )
    if cjk_chars >= 8:
        return True

    alpha_chars = [ch for ch in txt if ch.isalpha()]
    if not alpha_chars:
        return False

    ascii_alpha = sum(1 for ch in alpha_chars if ("a" <= ch.lower() <= "z"))
    ascii_ratio = ascii_alpha / max(1, len(alpha_chars))
    if ascii_ratio < 0.45:
        return True

    if _non_ascii_ratio(txt) <= 0.25:
        return False

    if re.search(r"\b(the|and|of|to|for|with|method|system|apparatus|device|invention)\b", txt, re.I):
        return False
    return True


def _extract_google_translate_text(payload: object) -> str:
    if not isinstance(payload, list) or not payload:
        return ""
    head = payload[0]
    if not isinstance(head, list):
        return ""
    parts: List[str] = []
    for row in head:
        if isinstance(row, list) and row:
            part = _clean(str(row[0] or ""))
            if part:
                parts.append(part)
    return " ".join(parts).strip()


def _translate_chunk_to_english(chunk: str) -> str:
    src = _clean(chunk)
    if not src:
        return ""
    url = (
        "https://translate.googleapis.com/translate_a/single?client=gtx"
        "&sl=auto&tl=en&dt=t&q="
        + quote(src, safe="")
    )
    try:
        with urlopen(url, timeout=_TRANSLATE_TIMEOUT_S) as resp:
            raw = resp.read().decode("utf-8", errors="ignore")
        translated = _extract_google_translate_text(json.loads(raw))
        return translated or src
    except (HTTPError, URLError, OSError, ValueError, json.JSONDecodeError):
        return src


def _translate_text_to_english(text: str) -> str:
    src = _clean_prior_art_abstract_text(text)
    if not src:
        return ""
    enabled = os.getenv("PRIOR_ART_TRANSLATE_TO_ENGLISH", "1").strip().lower()
    if enabled in {"0", "false", "no", "off"}:
        return src

    paras = [p.strip() for p in re.split(r"\n{2,}", src) if p and p.strip()]
    if not paras:
        paras = [src]

    out: List[str] = []
    for para in paras:
        if len(para) <= _TRANSLATE_MAX_CHARS:
            out.append(_translate_chunk_to_english(para))
            continue

        words = para.split()
        if not words:
            continue
        chunks: List[str] = []
        current: List[str] = []
        current_len = 0
        for w in words:
            add_len = len(w) + 1
            if current and (current_len + add_len) > _TRANSLATE_MAX_CHARS:
                chunks.append(" ".join(current))
                current = [w]
                current_len = len(w)
            else:
                current.append(w)
                current_len += add_len
        if current:
            chunks.append(" ".join(current))

        for chunk in chunks:
            out.append(_translate_chunk_to_english(chunk))

    translated = _clean_prior_art_abstract_text("\n\n".join(out))
    return translated or src


def _extract_prior_art_abstract_by_heading(lines: List[str]) -> str:
    if not lines:
        return ""
    headings = [h.lower() for h in _PRIOR_ART_ABSTRACT_HEADINGS]

    for i, line in enumerate(lines):
        x = _clean(line)
        if not x:
            continue
        low = x.lower()

        inline = ""
        if re.match(r"^\[?\s*57\s*\]?\s*abstract\b", low):
            m = re.match(r"^\[?\s*57\s*\]?\s*abstract\b\s*[:\-]?\s*(.*)$", x, re.I)
            inline = _clean(m.group(1) if m else "")
        else:
            for h in headings:
                m = re.match(rf"^\s*{re.escape(h)}\b\s*[:\-]?\s*(.*)$", x, re.I)
                if m:
                    inline = _clean(m.group(1))
                    break

        is_heading = bool(inline or low in headings or low in [f"{h}:" for h in headings] or re.match(r"^\[?\s*57\s*\]?\s*abstract\b\s*:?\s*$", low))
        if not is_heading:
            continue

        picked: List[str] = []
        wc = 0
        if inline:
            picked.append(inline)
            wc += len(inline.split())

        j = i + 1
        blank_streak = 0
        while j < len(lines):
            ln = _clean(lines[j])
            if not ln:
                blank_streak += 1
                # Layout extraction often introduces blank lines. Stop only after enough content.
                if wc >= 140 and blank_streak >= 2:
                    break
                j += 1
                continue
            blank_streak = 0
            if _is_prior_art_header_footer_noise(ln):
                if wc >= 120:
                    break
                j += 1
                continue
            if _looks_like_prior_art_heading(ln) and wc >= 90:
                break
            if _is_prior_art_metadata_line(ln):
                if wc < 20:
                    j += 1
                    continue
                if wc >= 120:
                    break
            picked.append(ln)
            wc += len(ln.split())
            if wc >= 750:
                break
            j += 1

        cand = _clean_prior_art_abstract_text("\n".join(picked))
        if len(cand.split()) >= 20:
            return cand
    return ""


def _extract_prior_art_abstract_fallback(lines: List[str]) -> str:
    if not lines:
        return ""

    blocks: List[str] = []
    cur: List[str] = []
    for ln in lines:
        x = _clean(ln)
        if not x:
            if cur:
                blocks.append(" ".join(cur))
                cur = []
            continue
        if _is_prior_art_metadata_line(x):
            if cur:
                blocks.append(" ".join(cur))
                cur = []
            continue
        cur.append(x)
    if cur:
        blocks.append(" ".join(cur))

    if not blocks:
        return ""

    def score(b: str) -> float:
        t = _clean(b)
        wc = len(t.split())
        if wc < 25:
            return -100.0
        s = 0.0
        if 45 <= wc <= 280:
            s += 5.0
        if re.search(r"\b(the present invention|discloses|relates to|provides|method|apparatus|system|device|implemented)\b", t, re.I):
            s += 2.0
        if re.search(r"\bclaim[s]?\b", t, re.I):
            s -= 2.5
        digit_ratio = sum(ch.isdigit() for ch in t) / max(1, len(t))
        if digit_ratio > 0.12:
            s -= 1.5
        return s

    best = max(blocks[:14], key=score)
    return _clean_prior_art_abstract_text(best)


def extract_prior_art_abstract_from_pdf(pdf_path: str) -> str:
    """Extract the most likely abstract text from a prior-art PDF."""
    lines = _read_prior_art_pdf_lines(pdf_path, max_pages=5)
    abstract = _extract_prior_art_abstract_by_heading(lines)
    if not abstract:
        abstract = _extract_prior_art_abstract_fallback(lines)

    if not abstract:
        full_text = read_pdf_text(pdf_path)
        abstract = _extract_prior_art_abstract_fallback(full_text.splitlines())

    abstract = _clean_prior_art_abstract_text(abstract)
    if abstract and _looks_non_english(abstract):
        abstract = _clean_prior_art_abstract_text(_translate_text_to_english(abstract))
    return _trim_abstract_without_mid_sentence_cut(abstract)


def _find_date(patterns: List[str], text: str) -> str:
    for pat in patterns:
        m = re.search(pat, text, re.I)
        if m:
            return _clean(m.group(1))
    return ""


DATE_RE = re.compile(r"\b(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})\b")
NUMERIC_DATE_RE = r"[0-9]{1,2}[./-][0-9]{1,2}[./-][0-9]{2,4}"


def _normalize_date(s: str) -> str:
    s = _clean(s)
    if not s:
        return ""
    s = s.replace(".", "/")
    s = s.replace("-", "/")
    m = DATE_RE.search(s)
    if not m:
        return ""
    d, mo, y = m.group(1), m.group(2), m.group(3)
    if len(y) == 2:
        y = "20" + y
    return f"{int(d):02d}/{int(mo):02d}/{y}"


def _canonical_docno(s: str) -> str:
    s = (s or "").upper()
    s = re.sub(r"\s+", "", s)
    s = re.sub(r"[()\-;:,]", "", s)
    return s


def _extract_hn_dispatch_date(text: str) -> str:
    txt = text or ""
    lines = [ln.strip() for ln in txt.splitlines() if ln and ln.strip()]

    patterns = [
        rf"date\s+of\s+dispatch(?:\s*/\s*email)?\s*[:\-]?\s*({NUMERIC_DATE_RE})",
        rf"dispatch\s+date\s*[:\-]?\s*({NUMERIC_DATE_RE})",
        rf"\bdispatch(?:ed)?\s+on\s*[:\-]?\s*({NUMERIC_DATE_RE})",
    ]
    for pat in patterns:
        m = re.search(pat, txt, re.I)
        if m:
            return _clean(m.group(1))

    # Common HN top-header format: "Date: dd/mm/yyyy"
    for ln in lines[:40]:
        low = ln.lower()
        if "hearing date" in low or "date & time" in low or "time" in low:
            continue
        if re.match(r"^date\s*[:\-]", low):
            m = re.search(NUMERIC_DATE_RE, ln)
            if m:
                return _clean(m.group(0))
        if re.fullmatch(NUMERIC_DATE_RE, ln):
            return _clean(ln)

    m = re.search(rf"hearing\s+notice\s+(?:is\s+)?(?:dated|date)\s*[:\-]?\s*({NUMERIC_DATE_RE})", txt, re.I)
    if m:
        return _clean(m.group(1))
    m = re.search(rf"\bnotice\s+dated\s*[:\-]?\s*({NUMERIC_DATE_RE})", txt, re.I)
    if m:
        return _clean(m.group(1))

    for ln in lines[:120]:
        low = ln.lower()
        if "dispatch" in low:
            m = re.search(NUMERIC_DATE_RE, ln)
            if m:
                return _clean(m.group(0))

    return ""


def _parse_prior_arts_from_text(text: str) -> List[PriorArt]:
    """Parse D1..Dn prior-arts from FER/HN text across IPO formats and de-duplicate."""
    lines = (text or "").splitlines()
    arts: List[PriorArt] = []
    i = 0
    dx_head = re.compile(r"^\s*(D\d+)\s*[:\-]?\s*(.*)$", re.I)
    stop_re = re.compile(
        r"^\s*(FORMAL\s+REQUIREMENT|REPLY\s+TO\s+OBJECTION|NOVELTY|INVENTIVE|NON[-\s]*PATENT|CLAIM|HEARING|NAME\s+OF\s+THE\s+CONTROLLER)\b",
        re.I,
    )

    while i < len(lines):
        raw = lines[i]
        m = dx_head.match(raw)
        if not m:
            i += 1
            continue

        label = _clean(m.group(1)).upper()
        rest = (m.group(2) or "").strip()

        block_parts = [rest] if rest else []
        j = i + 1
        while j < len(lines) and len(block_parts) < 6:
            nxt = lines[j]
            if dx_head.match(nxt):
                break
            if stop_re.match(nxt):
                break
            if _clean(nxt):
                block_parts.append(_clean(nxt))
            j += 1

        block = re.sub(r"\s+", " ", " ".join(block_parts)).strip()

        date = ""
        md = re.search(
            r"Publication\s*Date\s*[:\-]*\s*([0-9]{1,2}[/-][0-9]{1,2}[/-][0-9]{2,4})",
            block,
            re.I,
        )
        if md:
            date = _normalize_date(md.group(1))
        if not date:
            md = re.search(r"\(([0-9]{1,2}[/-][0-9]{1,2}[/-][0-9]{2,4})\)", block)
            if md:
                date = _normalize_date(md.group(1))
        if not date:
            date = _normalize_date(block)

        docno = block
        if date:
            mdt = re.search(r"[0-9]{1,2}[/-][0-9]{1,2}[/-][0-9]{2,4}", docno)
            if mdt:
                docno = docno[: mdt.start()].strip()
            docno = re.sub(r"Publication\s*Date\s*[:\-]*\s*$", "", docno, flags=re.I).strip()

        docno = re.sub(r"\bwhole\s+doc(?:ument)?\b.*$", "", docno, flags=re.I).strip()
        docno = docno.strip(" ;,")

        if docno:
            arts.append(PriorArt(label=label, docno=docno, date=date))

        i = j

    dedup: Dict[Tuple[str, str], PriorArt] = {}
    for pa in arts:
        key = (pa.label.upper(), _canonical_docno(pa.docno))
        prev = dedup.get(key)
        if prev is None:
            dedup[key] = pa
        else:
            if (not prev.date) and pa.date:
                dedup[key] = pa
            elif prev.date == pa.date and len(pa.docno) > len(prev.docno):
                dedup[key] = pa

    vals = list(dedup.values())
    vals.sort(key=lambda p: int(re.sub(r"\D", "", p.label) or "9999"))
    return vals


def _extract_disclosures(text: str, prior_arts: List[PriorArt]) -> Dict[str, str]:
    """Extract short disclosure strings for each Dx, robust to IPO phrasing."""
    t = text or ""
    out: Dict[str, str] = {}

    def clip(s: str) -> str:
        s = _clean(s)
        return s[:900]

    for pa in prior_arts or []:
        lab = pa.label.upper()

        m = re.search(
            rf"(?:Document\s+)?{re.escape(lab)}\s*[:\-]?\s*(?:{re.escape(pa.docno)}\s*)?(?:\([^)]+\)\s*)?.*?\b(discloses|describes|teaches|is\s+related\s+to)\b\s*(.*?)(?=(?:Document\s+)?D\d+\b|Therefore\b|Thus\b|Hence\b|In\s+view\b|\Z)",
            t,
            re.I | re.S,
        )
        if m:
            out[lab] = clip(m.group(2))
            continue

        m = re.search(
            rf"\b{re.escape(lab)}\b.*?\b(discloses|describes|teaches)\b\s*(.*?)(?=(?:Document\s+)?D\d+\b|Therefore\b|Thus\b|Hence\b|\Z)",
            t,
            re.I | re.S,
        )
        if m:
            out[lab] = clip(m.group(2))
            continue

        if pa.docno:
            doc_pat = re.escape(pa.docno)
            m = re.search(
                rf"{doc_pat}.*?\b(discloses|describes|teaches|is\s+related\s+to)\b\s*(.*?)(?=(?:Document\s+)?D\d+\b|{doc_pat}\b|Therefore\b|Thus\b|Hence\b|\Z)",
                t,
                re.I | re.S,
            )
            if m:
                out[lab] = clip(m.group(2))
                continue

    return out


def parse_case_meta_from_fer_or_hn(pdf_path: str) -> CaseMeta:
    text = read_pdf_text(pdf_path)
    meta = CaseMeta(prior_arts=[])

    # Application number
    for pat in [
        r"Indian\s+Patent\s+Application\s+No\s*[:\-]?\s*([0-9][0-9A-Z/\-]*)",
        r"Indian\s+Patent\s+Application\s+No\.?\s*[:\-]?\s*([0-9][0-9A-Z/\-]*)",
        r"Application\s+Number\s*[:\-]?\s*([0-9][0-9A-Z/\-]*)",
        r"Application\s*No\.?\s*[/:-]?\s*([0-9][0-9A-Z/\-]*)",
        r"POD/Application\s*No\s*/\s*([0-9][0-9A-Z/\-]*)",
    ]:
        m = re.search(pat, text, re.I)
        if m:
            meta.app_no = _clean(m.group(1))
            break

    # Filed on / Date of Filing
    meta.filed_on = _find_date(
        [
            r"Date\s+of\s+Filing\s*[:\-]?\s*([0-9]{1,2}[\-/][0-9]{1,2}[\-/][0-9]{2,4})",
            r"Filed\s*on\s*[:\-]?\s*([0-9]{1,2}[\-/][0-9]{1,2}[\-/][0-9]{2,4})",
        ],
        text,
    )

    # Applicant
    m = re.search(r"Name\s+of\s+the\s+Applicant\s*[:\-]?\s*(.+)", text, re.I)
    if not m:
        m = re.search(r"\bApplicant\s*[:\-]?\s*(.+)", text, re.I)
    if m:
        first = _clean(m.group(1))
        cont = ""
        after = text[m.end():]
        m2 = re.match(r"\s*\n\s*([^\n:]{4,})\n", after)
        if m2:
            cand = _clean(m2.group(1))
            if cand and not re.search(r"\b(controller|address|date|application|hearing|ref)\b", cand, re.I):
                cont = cand
        meta.applicant = _clean(" ".join([x for x in [first, cont] if x]))

    # Controller (expanded to cover IPO variants like:
    # "Saroj Kumar\nDeputy Controller of Patents & Designs"
    m = re.search(
        r"\n\s*([A-Za-z][A-Za-z .]{2,})\s*\n\s*"
        r"((?:Assistant|Deputy|Joint|Senior\s+Joint|Controller)\s+Controller\s+of\s+"
        r"(?:Patents?(?:\s*&\s*Designs)?|Patents?\s+and\s+Designs))\b",
        text,
        re.I,
    )
    if m:
        controller_name = _clean(m.group(1)).title()
        designation = _clean(m.group(2))
        meta.controller = f"{controller_name} ({designation})"
    else:
        m = re.search(r"Controller\s+Name\s*[:\-]?\s*(.+)", text, re.I)
        if m:
            meta.controller = _clean(m.group(1))
        else:
            meta.controller = ""

    # Agent
    m = re.search(r"To\s*\n\s*([A-Z][A-Z ]+NARASANI)", text, re.I)
    if m:
        meta.agents = _clean(m.group(1).title())
    m = re.search(r"Registerd\s+Address\s+For\s+Service\s*:?\s*([A-Z][^,\n]+NARASANI)", text, re.I)
    if m:
        meta.agents = _clean(m.group(1))

    # Dispatch dates
    meta.fer_dispatch_date = _find_date(
        [
            r"Date\s+of\s+Dispatch.*?:\s*([0-9]{1,2}[.\-/][0-9]{1,2}[.\-/][0-9]{2,4})",
            r"Date\s+of\s+Dispatch/Email.*?:\s*([0-9]{1,2}[.\-/][0-9]{1,2}[.\-/][0-9]{2,4})",
        ],
        text,
    )

    meta.hn_dispatch_date = _extract_hn_dispatch_date(text) or _find_date(
        [
            r"Date\s+of\s+Dispatch\s*:?\s*([0-9]{1,2}[.\-/][0-9]{1,2}[.\-/][0-9]{2,4})",
            r"Date\s+of\s+Dispatch/Email\s*:?\s*([0-9]{1,2}[.\-/][0-9]{1,2}[.\-/][0-9]{2,4})",
            r"hearing\s+notice\s+dated\s*([0-9]{1,2}[.\-/][0-9]{1,2}[.\-/][0-9]{2,4})",
            r"\bDate\s*[-:]\s*([0-9]{1,2}[.\-/][0-9]{1,2}[.\-/][0-9]{2,4})\b",
        ],
        text,
    )

    # Hearing date/time/mode
    m = re.search(r"Hearing\s+Location\s*:\s*(.+)", text, re.I)
    if m:
        meta.hearing_mode = _clean(m.group(1))
    m = re.search(
        r"Hearing\s+Date\s*&\s*Time\s*:\s*([0-9]{1,2}[\-/][0-9]{1,2}[\-/][0-9]{2,4})\s*/\s*([^\n]+)",
        text,
        re.I,
    )
    if m:
        meta.hearing_date = _clean(m.group(1))
        time_blob = _clean(m.group(2))
        tm = re.search(r"(\d{1,2}:\d{2})", time_blob)
        meta.hearing_time = tm.group(1) if tm else time_blob
        meta.hearing_duration = _duration_from_phrase(time_blob) or _duration_from_time_range(time_blob)

    if not meta.hearing_duration:
        meta.hearing_duration = _duration_from_phrase(text)
    if not meta.hearing_duration:
        m = re.search(r"Hearing\s+Duration\s*[:\-]\s*([^\n]+)", text, re.I)
        if m:
            meta.hearing_duration = _clean(m.group(1))
    if not meta.hearing_duration:
        m = re.search(r"\bDuration\s*[:\-]\s*([0-9]{1,3}\s*(?:minutes?|mins?|hours?|hrs?)(?:\s*[0-9]{1,2}\s*(?:minutes?|mins?))?)", text, re.I)
        if m:
            meta.hearing_duration = _clean(m.group(1))
    if not meta.hearing_duration:
        meta.hearing_duration = _duration_from_time_range(text)

    # FER date + reply date
    meta.fer_date = _find_date(
        [r"FER\s+dated\s*([0-9]{1,2}[\-/][0-9]{1,2}[\-/][0-9]{2,4})"],
        text,
    ) or meta.fer_dispatch_date

    meta.fer_reply_date = _find_date(
        [r"reply\s+of\s+the\s+applicant\s+dated\s*([0-9]{1,2}[\-/][0-9]{1,2}[\-/][0-9]{2,4})"],
        text,
    )

    # Prior arts + disclosures
    meta.prior_arts = _parse_prior_arts_from_text(text)

    disclosures = _extract_disclosures(text, meta.prior_arts or [])
    for lab, disc in disclosures.items():
        try:
            num = int(re.sub(r"\D", "", lab))
            setattr(meta, f"d{num}_disclosure", disc)
        except Exception:
            pass

    meta.d1_disclosure = getattr(meta, "d1_disclosure", "") or disclosures.get("D1", "")
    meta.d2_disclosure = getattr(meta, "d2_disclosure", "") or disclosures.get("D2", "")

    return meta

def build_prior_arts_list(meta: CaseMeta) -> str:
    lines: List[str] = []
    for pa in (meta.prior_arts or []):
        if pa.date:
            lines.append(f"{pa.label}: {pa.docno} ({pa.date})")
        else:
            lines.append(f"{pa.label}: {pa.docno}")
    return "\n".join(lines) if lines else ""


def parse_claims_from_specification(spec_text: str) -> Dict[int, str]:
    idx = spec_text.lower().find("claims")
    tail = spec_text[idx:] if idx != -1 else spec_text
    claims: Dict[int, str] = {}
    for m in re.finditer(r"(?ms)^\s*(\d{1,2})\.\s+(.*?)(?=^\s*\d{1,2}\.\s+|\Z)", tail):
        no = int(m.group(1))
        if 1 <= no <= 50:
            claims[no] = _clean(m.group(2))
    return claims


def parse_amended_claims(path: str) -> Dict[int, str]:
    text = read_text_any(path)
    if not text:
        return {}

    claims: Dict[int, str] = {}
    numbered = list(re.finditer(r"(?ms)^\s*(\d{1,2})\.\s+(.*?)(?=^\s*\d{1,2}\.\s+|\Z)", text))
    for m in numbered:
        no = int(m.group(1))
        if 1 <= no <= 200:
            claims[no] = _clean(m.group(2))

    if 1 not in claims:
        m = re.search(
            r"Claim\s*1\s+has\s+been\s+amended\s+to\s+recite\s*:\s*(.*?)(?=\n\s*TECHNICAL\s+ADVANCEMENT\s*:|\Z)",
            text,
            re.I | re.S,
        )
        if m:
            claims[1] = _clean(m.group(1))

    return claims


def extract_technical_advancement_from_spec(spec_text: str) -> str:
    txt = spec_text or ""
    if not txt.strip():
        return ""

    def norm_heading(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "")).strip().upper().rstrip(":")

    def is_noise_line(s: str) -> bool:
        x = re.sub(r"\s+", " ", (s or "")).strip()
        if not x:
            return True
        low = x.lower()
        if re.fullmatch(r"\d+\s*/\s*\d+", low):
            return True
        if re.fullmatch(r"page\s*\d+(\s*of\s*\d+)?", low):
            return True
        if re.fullmatch(r"date\s*[:\-]\s*\d{1,2}[./-]\d{1,2}[./-]\d{2,4}", low):
            return True
        if re.fullmatch(r"\d{1,2}[./-]\d{1,2}[./-]\d{2,4}", low):
            return True
        if "(cid:" in low:
            return True
        return False

    def strip_line_numbers_from_text(s: str) -> str:
        t = s or ""
        if not t:
            return ""
        t = re.sub(r"(?m)^\s*\d{1,3}\s*$", "", t)
        t = re.sub(r"(?m)^\s*\d{1,3}\s+(?=[A-Za-z\[\(])", "", t)

        def repl(m: re.Match) -> str:
            try:
                n = int(m.group(1))
            except Exception:
                return m.group(0)
            if 1 <= n <= 400 and n % 5 == 0:
                return " "
            return m.group(0)

        t = re.sub(r"(?<=[A-Za-z\)])\s+(\d{1,3})\s+(?=[A-Za-z\(\[])", repl, t)
        t = re.sub(r"[ \t]{2,}", " ", t)
        t = re.sub(r"\n{3,}", "\n\n", t)
        return t.strip()

    lines = [ln.rstrip() for ln in txt.splitlines()]
    start_idx = -1
    detail_heads = [
        "DETAILED DESCRIPTION OF THE INVENTION",
        "DETAILED DESCRIPTION OF INVENTION",
        "DETAILED DESCRIPTION",
    ]
    for i, ln in enumerate(lines):
        h = norm_heading(ln)
        if h in detail_heads:
            start_idx = i
            break
    if start_idx < 0:
        numbered = re.findall(r"(?ms)(\[\d{4}\].*?)(?=\n\s*\[\d{4}\]|\Z)", txt)
        if not numbered:
            return ""
        cleaned_num = [strip_line_numbers_from_text(re.sub(r"\s+", " ", p).strip()) for p in numbered if p and p.strip()]
        preferred = []
        for p in cleaned_num:
            mnum = re.match(r"\[(\d{4})\]", p)
            if mnum and int(mnum.group(1)) >= 30:
                preferred.append(p)
        pick = preferred or cleaned_num
        return strip_line_numbers_from_text("\n\n".join(pick[:4]).strip())

    end_heads = {
        "CLAIMS",
        "ABSTRACT",
        "WE CLAIM",
        "WHAT IS CLAIMED IS",
    }
    end_idx = len(lines)
    for j in range(start_idx + 1, len(lines)):
        h = norm_heading(lines[j])
        if h in end_heads:
            end_idx = j
            break

    section_lines = []
    for ln in lines[start_idx + 1 : end_idx]:
        if is_noise_line(ln):
            continue
        norm_ln = re.sub(r"[ \t]+", " ", ln).strip()
        norm_ln = re.sub(r"^\d{1,3}\s+(?=[A-Za-z\[\(])", "", norm_ln)
        if norm_ln:
            section_lines.append(norm_ln)
    section = "\n".join(section_lines).strip()
    if not section:
        return ""

    numbered_paras = re.findall(r"(?ms)(\[\d{4}\].*?)(?=\n\s*\[\d{4}\]|\Z)", section)
    if numbered_paras:
        cleaned = [strip_line_numbers_from_text(re.sub(r"\s+", " ", p).strip()) for p in numbered_paras if p and p.strip()]
        return strip_line_numbers_from_text("\n\n".join(cleaned[:4]).strip())

    para_blocks = [strip_line_numbers_from_text(re.sub(r"\s+", " ", p).strip()) for p in re.split(r"\n\s*\n+", section) if p and p.strip()]
    return strip_line_numbers_from_text("\n\n".join(para_blocks[:4]).strip())


def extract_figure_descriptions_from_spec(spec_text: str) -> Dict[int, str]:
    txt = spec_text or ""
    out: Dict[int, str] = {}
    lines = [re.sub(r"\s+", " ", ln).strip() for ln in txt.splitlines()]
    fig_re = re.compile(r"\bFIG\.?\s*(\d+)(?:[A-Z])?\b\s*(?:is|illustrates|shows|depicts|represents)?\s*[:\-]?\s*(.*)$", re.I)
    for ln in lines:
        m = fig_re.search(ln)
        if not m:
            continue
        num = int(m.group(1))
        desc = (m.group(2) or "").strip()
        if not desc:
            continue
        desc = desc.strip().strip(";").strip()
        if num not in out or (len(desc) > len(out[num])):
            out[num] = desc[:300]
    return out
